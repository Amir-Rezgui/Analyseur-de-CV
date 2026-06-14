"""
Report Generator — Professional PDF and Word reports.
Clean, corporate-quality output suitable for real-world presentations.
Updated to show tier classification, synonym/related matches.
"""

import os
from datetime import datetime
from typing import List

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Flowable, PageBreak
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.scorer import ScoringResult


# ─── PDF Custom Flowables ────────────────────────────────────────────

class ColorBar(Flowable):
    def __init__(self, width, height, fill_ratio, fill_color, bg_color=None):
        Flowable.__init__(self)
        self.width = width
        self.height = height
        self.ratio = min(1.0, max(0.0, fill_ratio))
        self.fill_color = fill_color
        self.bg_color = bg_color or colors.HexColor("#f1f5f9")

    def draw(self):
        # background
        self.canv.setFillColor(self.bg_color)
        self.canv.roundRect(0, 0, self.width, self.height, self.height/2, fill=1, stroke=0)
        # fill
        if self.ratio > 0:
            fw = max(self.height, self.width * self.ratio)
            self.canv.setFillColor(self.fill_color)
            self.canv.roundRect(0, 0, fw, self.height, self.height/2, fill=1, stroke=0)


class ScoreCircle(Flowable):
    def __init__(self, score, size=90):
        Flowable.__init__(self)
        self.score = score
        self.size = size
        self.width = size
        self.height = size

    def draw(self):
        c = self.canv
        cx = cy = self.size / 2
        r = self.size / 2 - 5
        color = self._color()

        c.setStrokeColor(colors.HexColor("#e8ecf0"))
        c.setLineWidth(6)
        c.circle(cx, cy, r, fill=0)

        c.setStrokeColor(color)
        c.setLineWidth(6)
        extent = (self.score / 100) * 360
        c.arc(cx - r, cy - r, cx + r, cy + r, 90, extent)

        c.setFillColor(colors.HexColor("#111827"))
        c.setFont("Helvetica-Bold", 20)
        txt = f"{self.score:.0f}%"
        tw = c.stringWidth(txt, "Helvetica-Bold", 20)
        c.drawString(cx - tw/2, cy - 6, txt)

        c.setFillColor(colors.HexColor("#6b7280"))
        c.setFont("Helvetica", 7)
        lbl = "SCORE"
        lw = c.stringWidth(lbl, "Helvetica", 7)
        c.drawString(cx - lw/2, cy - 18, lbl)

    def _color(self):
        if self.score >= 82: return colors.HexColor("#059669")
        if self.score >= 65: return colors.HexColor("#0284c7")
        if self.score >= 45: return colors.HexColor("#d97706")
        return colors.HexColor("#dc2626")


# ─── Main Generator ─────────────────────────────────────────────────

class ReportGenerator:

    # Professional charcoal + emerald palette
    C_BRAND   = colors.HexColor("#10b981")  # emerald accent
    C_DARK    = colors.HexColor("#111827")  # near-black for headings
    C_MEDIUM  = colors.HexColor("#374151")  # body text
    C_MUTED   = colors.HexColor("#6b7280")  # secondary text
    C_BG      = colors.HexColor("#f9fafb")  # light background
    C_BG2     = colors.HexColor("#f3f4f6")  # alternate background
    C_SUCCESS = colors.HexColor("#059669")  # emerald success
    C_WARNING = colors.HexColor("#d97706")  # amber
    C_DANGER  = colors.HexColor("#dc2626")  # red
    C_INFO    = colors.HexColor("#0284c7")  # sky blue
    C_TEAL    = colors.HexColor("#0d9488")  # teal accent
    C_WHITE   = colors.white
    C_BORDER  = colors.HexColor("#e5e7eb")

    PAGE_W, PAGE_H = A4

    TIER_LABELS = {
        "critical": "Critique",
        "important": "Important",
        "noise": "Secondaire",
    }

    # ── PDF ─────────────────────────────────────────────────────────

    def generate_pdf(self, result: ScoringResult, output_path: str,
                     cv_name: str = "CV", job_name: str = "Poste") -> str:
        doc = SimpleDocTemplate(
            output_path, pagesize=A4,
            rightMargin=2.2*cm, leftMargin=2.2*cm,
            topMargin=1.8*cm, bottomMargin=2*cm,
        )
        W = self.PAGE_W - 4.4*cm
        styles = getSampleStyleSheet()

        def style(name, **kw):
            return ParagraphStyle(name, parent=styles['Normal'], **kw)

        S = {
            'title':    style('T', fontName='Helvetica-Bold', fontSize=24,
                              textColor=self.C_DARK, leading=28, spaceAfter=4),
            'subtitle': style('Sub', fontName='Helvetica', fontSize=10,
                              textColor=self.C_MUTED, spaceAfter=18),
            'h2':       style('H2', fontName='Helvetica-Bold', fontSize=13,
                              textColor=self.C_DARK, spaceBefore=22, spaceAfter=10, leading=17),
            'label':    style('Lbl', fontName='Helvetica', fontSize=8.5,
                              textColor=self.C_MUTED, spaceAfter=2),
            'value':    style('Val', fontName='Helvetica-Bold', fontSize=11,
                              textColor=self.C_DARK),
            'body':     style('Bd', fontName='Helvetica', fontSize=9.5,
                              textColor=self.C_MEDIUM, spaceAfter=4, leading=14),
            'small':    style('Sm', fontName='Helvetica', fontSize=8.5,
                              textColor=self.C_MUTED, spaceAfter=3, leading=12),
            'bullet':   style('Bul', fontName='Helvetica', fontSize=9.5,
                              textColor=self.C_MEDIUM, spaceAfter=5, leading=14,
                              leftIndent=14),
            'cat_pct':  style('CP', fontName='Helvetica-Bold', fontSize=11,
                              alignment=TA_RIGHT),
            'footer':   style('Ft', fontName='Helvetica', fontSize=7.5,
                              textColor=self.C_MUTED, alignment=TA_CENTER, spaceBefore=18),
            'tier_hdr': style('TH', fontName='Helvetica-Bold', fontSize=10,
                              textColor=self.C_MEDIUM, spaceBefore=14, spaceAfter=6),
        }

        els = []

        # ── Header stripe — dual-tone professional ──
        els.append(ColorBar(W, 5, 1.0, self.C_DARK))
        els.append(Spacer(1, 20))

        els.append(Paragraph("Rapport d\u2019Évaluation de Compatibilité", S['title']))
        els.append(Paragraph(
            f"CV Fit Analyzer  ·  {datetime.now().strftime('%d %B %Y, %H:%M')}",
            S['subtitle']))

        # ── Score overview ──
        gauge = ScoreCircle(result.overall_score, size=90)
        score_color = self._score_color(result.overall_score)

        kv = [
            ("Score Global",              f"{result.overall_score:.1f} / 100"),
            ("Niveau de compatibilité",   result.classification),
            ("Compétences matchées",      f"{result.total_matched} sur {result.total_required} requises"),
        ]
        if result.total_synonyms > 0:
            kv.append(("Compétences équivalentes", str(result.total_synonyms)))
        if result.total_related > 0:
            kv.append(("Compétences proches", str(result.total_related)))
            
        kv_data = [[
            Paragraph(f'<font color="#6b7280">{k}</font>', S['label']),
            Paragraph(f'<font color="#111827"><b>{v}</b></font>', S['body']),
        ] for k, v in kv]

        kv_tbl = Table(kv_data, colWidths=[5*cm, 6*cm])
        kv_tbl.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
        ]))

        overview_inner = Table([[gauge, kv_tbl]], colWidths=[3.2*cm, W-3.2*cm])
        overview_inner.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING', (1,0), (1,0), 20),
            ('LEFTPADDING', (0,0), (0,0), 0),
        ]))

        overview_box = Table([[overview_inner]], colWidths=[W],
            style=TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), self.C_BG),
                ('ROUNDEDCORNERS', [8,8,8,8]),
                ('TOPPADDING', (0,0), (-1,-1), 18),
                ('BOTTOMPADDING', (0,0), (-1,-1), 18),
                ('LEFTPADDING', (0,0), (-1,-1), 18),
                ('RIGHTPADDING', (0,0), (-1,-1), 18),
            ]))
        els.append(overview_box)
        els.append(Spacer(1, 6))

        # ── Category breakdown by tier ──
        els.append(Paragraph("Détail par Catégorie", S['h2']))

        # Group by tier
        tier_order = ["critical", "important", "noise"]
        tier_titles = {
            "critical": "Compétences Techniques Critiques",
            "important": "Compétences Importantes",
            "noise": "Soft Skills (impact minimal sur le score)",
        }

        for tier in tier_order:
            tier_cats = [cs for cs in result.category_scores if cs.tier == tier]
            if not tier_cats:
                continue
            
            els.append(Paragraph(tier_titles.get(tier, tier), S['tier_hdr']))

            for cs in tier_cats:
                cat_color = self._ratio_color(cs.match_ratio)
                pct = f"{cs.match_ratio*100:.0f}%"

                hdr = Table([[
                    Paragraph(f'<b>{cs.category_label}</b>', S['body']),
                    Paragraph(
                        f'<font color="{cat_color.hexval()}">{pct}</font>',
                        ParagraphStyle('_cp', fontName='Helvetica-Bold', fontSize=10,
                                       textColor=cat_color, alignment=TA_RIGHT)
                    ),
                ]], colWidths=[W-2.5*cm, 2.5*cm])
                hdr.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 2),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 2),
                ]))
                els.append(hdr)
                els.append(Spacer(1, 4))
                els.append(ColorBar(W, 6, cs.match_ratio, cat_color, self.C_BG2))
                els.append(Spacer(1, 5))

                parts = []
                for sk in cs.matched_skills:
                    parts.append(f'<font color="#059669">✓ {sk}</font>')
                for syn in cs.synonym_skills:
                    parts.append(f'<font color="#0284c7">≈ {syn["job"]} (→{syn["cv"]})</font>')
                for rel in cs.related_skills:
                    parts.append(f'<font color="#d97706">↔ {rel["job"]} (→{rel["cv"]})</font>')
                for sk in cs.missing_skills:
                    parts.append(f'<font color="#dc2626">✗ {sk}</font>')
                for sk in cs.extra_skills[:3]:
                    parts.append(f'<font color="#0d9488">+ {sk}</font>')

                if parts:
                    els.append(Paragraph('   '.join(parts), S['small']))
                els.append(Spacer(1, 10))

        # ── Insights ──
        for title, items in [
            ("Points Forts", result.strengths),
            ("Points à Améliorer", result.weaknesses),
            ("Recommandations", result.recommendations),
        ]:
            if items:
                els.append(Paragraph(title, S['h2']))
                for item in items:
                    els.append(Paragraph(f"<bullet>•</bullet> {item}", S['bullet']))

        # ── Footer ──
        els.append(Spacer(1, 28))
        els.append(HRFlowable(width="100%", color=self.C_BORDER, thickness=1))
        els.append(Spacer(1, 6))
        els.append(Paragraph(
            "CV Fit Analyzer  ·  Rapport d'évaluation de compatibilité CV ↔ Fiche de poste",
            S['footer']))

        doc.build(els)
        return output_path

    # ── Word ────────────────────────────────────────────────────────

    def generate_docx(self, result: ScoringResult, output_path: str,
                      cv_name: str = "CV", job_name: str = "Poste") -> str:
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.8)
            section.right_margin  = Cm(2.8)

        # Default Normal style
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.space_before = Pt(0)

        # Heading styles
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(12)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        h1.paragraph_format.space_before = Pt(16)
        h1.paragraph_format.space_after  = Pt(6)

        h2 = doc.styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(10)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after  = Pt(4)

        # ── Title block ──
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(4)
        run = title_p.add_run("Rapport d\u2019Évaluation de Compatibilité")
        run.font.name = 'Calibri'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_after = Pt(14)
        gen_time = datetime.now().strftime('%d %B %Y à %H:%M')
        r = date_p.add_run(f"CV Fit Analyzer  ·  Généré le {gen_time}")
        r.font.name = 'Calibri'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

        self._add_hline(doc)

        # ── Score Overview ──
        doc.add_heading('Score Global', level=1)

        hdrs = ['Score', 'Niveau', 'Matchées', 'Requises']
        vals = [
            f"{result.overall_score:.1f}%",
            result.classification,
            str(result.total_matched),
            str(result.total_required),
        ]
        
        t = doc.add_table(rows=2, cols=len(hdrs))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, h in enumerate(hdrs):
            c = t.rows[0].cells[i]
            c.paragraphs[0].clear()
            run = c.paragraphs[0].add_run(h)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(8)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, v in enumerate(vals):
            c = t.rows[1].cells[i]
            c.paragraphs[0].clear()
            run = c.paragraphs[0].add_run(v)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(14)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._style_table(t)
        
        # Extra info
        if result.total_synonyms > 0 or result.total_related > 0:
            info_p = doc.add_paragraph()
            info_p.paragraph_format.space_before = Pt(6)
            parts = []
            if result.total_synonyms > 0:
                parts.append(f"{result.total_synonyms} compétence(s) équivalente(s)")
            if result.total_related > 0:
                parts.append(f"{result.total_related} compétence(s) proche(s)")
            r = info_p.add_run("Matching intelligent : " + " · ".join(parts))
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

        doc.add_paragraph()

        # ── Category Breakdown by Tier ──
        tier_order = ["critical", "important", "noise"]
        tier_titles = {
            "critical": "Compétences Critiques",
            "important": "Compétences Importantes",
            "noise": "Soft Skills (faible impact)",
        }

        for tier in tier_order:
            tier_cats = [cs for cs in result.category_scores if cs.tier == tier]
            if not tier_cats:
                continue
            
            doc.add_heading(tier_titles.get(tier, tier), level=1)

            cat_t = doc.add_table(rows=1 + len(tier_cats), cols=4)
            cat_hdrs = ['Catégorie', 'Match %', 'Matchées/Requises', 'Compétences Manquantes']
            for i, h in enumerate(cat_hdrs):
                c = cat_t.rows[0].cells[i]
                c.paragraphs[0].clear()
                run = c.paragraphs[0].add_run(h)
                run.font.name  = 'Calibri'
                run.font.size  = Pt(8)
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

            for idx, cs in enumerate(tier_cats, 1):
                row = cat_t.rows[idx]
                
                # Build missing list with synonym/related info
                missing_display = cs.missing_skills[:4]
                missing_str = ", ".join(missing_display)
                if len(cs.missing_skills) > 4:
                    missing_str += f" (+{len(cs.missing_skills)-4})"
                if not missing_str:
                    missing_str = "—"
                
                matched_total = len(cs.matched_skills) + len(cs.synonym_skills)
                required_total = matched_total + len(cs.related_skills) + len(cs.missing_skills)
                
                data = [
                    cs.category_label,
                    f"{cs.match_ratio*100:.0f}%",
                    f"{matched_total}/{required_total}",
                    missing_str,
                ]
                for i, val in enumerate(data):
                    c = row.cells[i]
                    c.paragraphs[0].clear()
                    run = c.paragraphs[0].add_run(val)
                    run.font.name  = 'Calibri'
                    run.font.size  = Pt(9)
                    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

            self._style_table(cat_t)
            doc.add_paragraph()

        # ── Insights ──
        for title, items in [
            ('Points Forts',        result.strengths),
            ("Points à Améliorer",  result.weaknesses),
            ('Recommandations',     result.recommendations),
        ]:
            if not items:
                continue
            doc.add_heading(title, level=1)
            for item in items:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(item)
                run.font.name  = 'Calibri'
                run.font.size  = Pt(10)
                run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

        # ── Footer ──
        doc.add_paragraph()
        self._add_hline(doc)
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(6)
        r = fp.add_run("CV Fit Analyzer  ·  Rapport d'évaluation de compatibilité CV ↔ Fiche de poste")
        r.font.name  = 'Calibri'
        r.font.size  = Pt(8)
        r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

        doc.save(output_path)
        return output_path

    # ── Helpers ─────────────────────────────────────────────────────

    def _score_color(self, score: float) -> colors.Color:
        if score >= 82: return self.C_SUCCESS
        if score >= 65: return self.C_INFO
        if score >= 45: return self.C_WARNING
        return self.C_DANGER

    def _ratio_color(self, ratio: float) -> colors.Color:
        if ratio >= 0.8: return self.C_SUCCESS
        if ratio >= 0.6: return self.C_INFO
        if ratio >= 0.4: return self.C_WARNING
        return self.C_DANGER

    def _style_table(self, table):
        """Apply clean border styling to a Word table."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Table width — auto
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '0')
        tblW.set(qn('w:type'), 'auto')
        tblPr.append(tblW)

        # Borders
        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   'single')
            b.set(qn('w:sz'),    '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'E5E7EB')
            tblBorders.append(b)
        tblPr.append(tblBorders)

        # Cell padding
        tblCellMar = OxmlElement('w:tblCellMar')
        for side in ('top','left','bottom','right'):
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'),    '80')
            m.set(qn('w:type'), 'dxa')
            tblCellMar.append(m)
        tblPr.append(tblCellMar)

    def _add_hline(self, doc):
        """Add a subtle horizontal rule paragraph."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'E5E7EB')
        pBdr.append(bot)
        pPr.append(pBdr)
